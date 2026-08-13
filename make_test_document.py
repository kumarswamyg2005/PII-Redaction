#!/usr/bin/env python3
"""
Generate a prospectus-shaped .docx the tool has never seen, with exact ground truth.

Every claim about generalisation made from the supplied prospectus is weak,
because the tool was developed while looking at that file. This script builds a
different document — different people, different companies, different addresses,
different identifiers — and records what it planted while it writes.

That matters for two reasons:

* **The annotations are exact by construction.** Nothing is transcribed by hand,
  so there is no annotation gap to argue about; a miss is a real miss and a
  false positive is real over-redaction.
* **Nothing here appears in the source prospectus.** Names, companies and
  addresses are drawn from pools that share no entries with it, so a tool that
  had memorised the original would score zero.

Run:  python make_test_document.py
      python -m pii_redaction synthetic_test.docx out.docx \\
             --report synthetic_report.md --ground-truth synthetic_ground_truth.json
"""

from __future__ import annotations

import json
import random
from pathlib import Path

from docx import Document
from docx.shared import Pt

from pii_redaction.detection.validators import verhoeff_check_digit

OUTPUT_DOCX = Path("synthetic_test.docx")
OUTPUT_TRUTH = Path("synthetic_ground_truth.json")
SEED = 20260814

# Deliberately disjoint from anything in the supplied prospectus.
FORENAMES = ["Bhavesh", "Charulata", "Devendra", "Ekta", "Farhan", "Gauri",
             "Harshad", "Indrani", "Jatin", "Komal", "Lalit", "Mrinalini"]
SURNAMES = ["Adhikari", "Bhagat", "Chitnis", "Dandekar", "Ekbote", "Fernandes",
            "Gokhale", "Hublikar", "Inamdar", "Jagtap", "Karnik", "Limaye"]
COMPANY_STEMS = ["Vantara", "Suryakiran", "Nilgiri", "Chandrapur", "Ratnagiri",
                 "Bhimashankar", "Kalyani", "Girivan"]
COMPANY_TAILS = ["Alloys Limited", "Textiles Private Limited", "Logistics Limited",
                 "Chemicals Private Limited", "Infrastructure Limited"]
STREETS = ["Survey No. 44/2, Ambegaon Khurd", "Plot 17, Sector 9, Kharghar",
           "Gat No. 88, Talegaon Dabhade", "Unit 3, Rajiv Gandhi Infotech Park"]
LOCALITIES = ["Hadapsar", "Bavdhan", "Kothrud", "Wagholi", "Undri"]


def _aadhaar(rng: random.Random) -> str:
    body = str(rng.randint(2, 9)) + "".join(str(rng.randint(0, 9)) for _ in range(10))
    number = body + verhoeff_check_digit(body)
    return f"{number[:4]} {number[4:8]} {number[8:]}"


def _pan(rng: random.Random) -> str:
    letters = lambda n: "".join(rng.choice("ABCDEFGHIJKLMNOPQRSTUVWXYZ") for _ in range(n))
    return f"{letters(3)}P{letters(1)}{rng.randint(1000, 9999)}{letters(1)}"


def main() -> int:
    rng = random.Random(SEED)
    document = Document()
    document.styles["Normal"].font.size = Pt(10)

    planted: list[dict] = []

    def say(text: str, pii: list[tuple[str, str]] | None = None, style: str | None = None) -> None:
        """Write a paragraph and record the PII values it contains."""
        document.add_paragraph(text, style=style)
        for entity_type, value in pii or []:
            planted.append({"type": entity_type, "text": value})

    people = [f"{rng.choice(FORENAMES)} {rng.choice(SURNAMES)}" for _ in range(6)]
    people = list(dict.fromkeys(people))
    companies = [
        f"{rng.choice(COMPANY_STEMS)} {rng.choice(COMPANY_TAILS)}" for _ in range(5)
    ]
    companies = list(dict.fromkeys(companies))
    issuer = companies[0]

    say("DRAFT RED HERRING PROSPECTUS", style="Title")
    say(f"{issuer} was incorporated under the Companies Act, 2013 and is "
        f"registered with the Registrar of Companies, Pune.",
        [("ORGANIZATION", issuer)])

    say("Definitions and Abbreviations")
    table = document.add_table(rows=1, cols=2)
    table.rows[0].cells[0].text = "Term"
    table.rows[0].cells[1].text = "Description"
    for term, description in [
        ("ASBA", "Application Supported by Blocked Amount, the mechanism by which "
                 "a Bidder authorises blocking of funds in the bank account"),
        ("Working Day", "All days on which commercial banks in Mumbai are open for "
                        "business, excluding Sundays and public holidays"),
        ("Offer Price", "The final price at which Equity Shares will be Allotted "
                        "to successful Bidders in terms of this document"),
    ]:
        row = table.add_row()
        row.cells[0].text = term
        row.cells[1].text = description

    say("Promoters and Key Managerial Personnel")
    for person in people[:3]:
        aadhaar, pan = _aadhaar(rng), _pan(rng)
        email = f"{person.split()[0].lower()}.{person.split()[1].lower()}@example-corp.com"
        phone = f"+91 {rng.randint(70000, 99999)} {rng.randint(10000, 99999)}"
        say(f"{person} holds an Aadhaar number {aadhaar} and PAN {pan}. "
            f"Contact: {email}, Telephone: {phone}.",
            [("PERSON", person), ("IN_AADHAAR", aadhaar), ("IN_PAN", pan),
             ("EMAIL_ADDRESS", email), ("PHONE_NUMBER", phone)])

    say("Registered and Corporate Offices")
    for company in companies[1:4]:
        address = (f"{rng.choice(STREETS)}, {rng.choice(LOCALITIES)}, "
                   f"Pune - 4{rng.randint(11000, 11999)}, Maharashtra, India")
        say(f"{company}, Registered Office: {address}",
            [("ORGANIZATION", company), ("LOCATION", address)])

    say("Reference numbers below identify filings, not people, and are retained.")
    say(f"Order No. {rng.randint(100000, 999999)} and Ticket No. "
        f"{rng.randint(100000, 999999)} relate to the escrow instruction.")

    say("The Offer is being made through the Book Building Process in terms of "
        "the SEBI ICDR Regulations. ASBA Bidders shall approach the SCSBs. The "
        "Equity Shares are proposed to be listed on BSE and NSE.")

    document.save(OUTPUT_DOCX)

    # Locate each planted value in the finished document to produce offsets.
    from pii_redaction.documents.docx_reader import joined_text, read_paragraphs

    text = joined_text(read_paragraphs(OUTPUT_DOCX))
    annotations, seen = [], set()
    for item in planted:
        start = text.find(item["text"])
        while start != -1:
            key = (start, start + len(item["text"]))
            if key not in seen:
                seen.add(key)
                annotations.append({
                    "type": item["type"], "start": key[0], "end": key[1],
                    "text": item["text"],
                })
                break
            start = text.find(item["text"], start + 1)

    annotations.sort(key=lambda a: a["start"])
    OUTPUT_TRUTH.write_text(json.dumps({
        "document": OUTPUT_DOCX.name,
        "region": {"start": 0, "end": len(text)},
        "note": "Generated document; annotations are exact by construction. "
                "The whole document is the scored region.",
        "annotations": annotations,
    }, indent=2))

    print(f"wrote {OUTPUT_DOCX} ({len(text)} chars)")
    print(f"wrote {OUTPUT_TRUTH} ({len(annotations)} annotations)")
    by_type: dict[str, int] = {}
    for annotation in annotations:
        by_type[annotation["type"]] = by_type.get(annotation["type"], 0) + 1
    for entity_type, count in sorted(by_type.items()):
        print(f"  {entity_type:16} {count}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
