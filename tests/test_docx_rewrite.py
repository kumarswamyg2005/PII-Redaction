"""
Rewriting a .docx in place.

These tests build small documents rather than using the prospectus, so they
stay fast and assert one behaviour each. Every case here corresponds to a real
defect found while building the tool.
"""

import zipfile

import pytest
from docx import Document

from pii_redaction.documents.docx_writer import DocxRedactor
from pii_redaction.surrogates.matcher import EntityMatcher


@pytest.fixture
def split_run_document(tmp_path):
    """A paragraph whose entity is spread over several runs, as Word does."""
    path = tmp_path / "split.docx"
    document = Document()
    paragraph = document.add_paragraph()
    for fragment, bold in [
        ("The promoter ", False), ("Everest", True), (" ", True),
        ("Family", True), (" ", True), ("Trust", True), (" holds shares.", False),
    ]:
        run = paragraph.add_run(fragment)
        run.bold = bold
    document.save(path)
    return path


def test_entity_split_across_runs_is_replaced_once(split_run_document, tmp_path):
    out = tmp_path / "out.docx"
    stats = DocxRedactor({"everest family trust": "Acme Holdings Limited"}).redact(
        split_run_document, out
    )
    text = "\n".join(p.text for p in Document(out).paragraphs)
    assert "Redwood Family Trust" not in text
    assert text.count("Acme Holdings Limited") == 1
    assert stats.total == 1


def test_formatting_outside_the_match_is_preserved(split_run_document, tmp_path):
    out = tmp_path / "out.docx"
    DocxRedactor({"everest family trust": "Acme Holdings Limited"}).redact(
        split_run_document, out
    )
    paragraph = Document(out).paragraphs[0]
    assert paragraph.runs[0].text == "The promoter "
    assert not paragraph.runs[0].bold
    assert paragraph.runs[-1].text == " holds shares."


def test_longest_entity_wins(tmp_path):
    source, out = tmp_path / "a.docx", tmp_path / "b.docx"
    document = Document()
    document.add_paragraph("Example Cables Limited and Example Cables Private Limited")
    document.save(source)

    DocxRedactor({
        "example cables": "SHORT",
        "example cables limited": "Paramount Limited",
    }).redact(source, out)
    text = Document(out).paragraphs[0].text
    assert "Paramount Limited" in text
    assert "Paramount Limited Limited" not in text


def test_hyperlink_field_targets_are_rewritten(tmp_path):
    """A link can expose a real address even when the visible text is clean."""
    source, out = tmp_path / "a.docx", tmp_path / "b.docx"
    document = Document()
    document.add_paragraph("Contact us")
    document.save(source)

    # Inject a HYPERLINK field instruction the way Word stores one.
    with zipfile.ZipFile(source) as archive:
        payload = {n: archive.read(n) for n in archive.namelist()}
    xml = payload["word/document.xml"].decode()
    xml = xml.replace(
        "</w:body>",
        '<w:p><w:r><w:instrText xml:space="preserve"> HYPERLINK '
        '"mailto:real.person@example-corp.com" </w:instrText></w:r></w:p></w:body>',
    )
    payload["word/document.xml"] = xml.encode()
    with zipfile.ZipFile(source, "w") as archive:
        for name, data in payload.items():
            archive.writestr(name, data)

    DocxRedactor({"real.person@example-corp.com": "fake@example.com"}).redact(source, out)
    with zipfile.ZipFile(out) as archive:
        result = archive.read("word/document.xml").decode()
    assert "real.person@example-corp.com" not in result
    assert "fake@example.com" in result


def test_document_structure_is_unchanged(split_run_document, tmp_path):
    out = tmp_path / "out.docx"
    DocxRedactor({"everest family trust": "Acme Holdings Limited"}).redact(
        split_run_document, out
    )
    with zipfile.ZipFile(split_run_document) as a, zipfile.ZipFile(out) as b:
        assert a.namelist() == b.namelist()


class TestMatcher:
    def test_does_not_fire_inside_an_email_local_part(self):
        """Replacing a fragment of an address breaks it and leaks the domain."""
        matcher = EntityMatcher({"ksh": "Acme"})
        text, count = matcher.apply("write to ksh@in.mpms.example.com today")
        assert count == 0
        assert "ksh@in.mpms.example.com" in text

    def test_replaces_a_whole_address(self):
        matcher = EntityMatcher({"ksh@in.mpms.example.com": "fake@example.com"})
        text, count = matcher.apply("write to ksh@in.mpms.example.com today")
        assert count == 1
        assert "fake@example.com" in text

    def test_respects_word_boundaries(self):
        matcher = EntityMatcher({"us": "THEM"})
        text, _ = matcher.apply("business is business")
        assert text == "business is business"

    def test_matches_across_a_line_break(self):
        matcher = EntityMatcher({"acme holdings": "Fake Corp"})
        text, count = matcher.apply("owned by Acme\nHoldings today")
        assert count == 1
        assert "Fake Corp" in text
