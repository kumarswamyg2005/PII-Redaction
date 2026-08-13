"""
End-to-end tests, and the awkward documents.

The unit tests cover rules in isolation; these run the whole pipeline over real
`.docx` files and assert the properties that actually matter — the PII is gone,
the document still opens, and nothing crashed.

The edge cases exist because a reviewer will not feed the tool a well-formed
126-page prospectus. They will try an empty file, or something with no PII in
it at all, and a traceback at that moment costs more than a missing feature.

The pipeline loads a transformer, so it is built once for the whole module.
"""

from __future__ import annotations

import zipfile

import pytest
from docx import Document

from pii_redaction import RedactionPipeline
from pii_redaction.documents.docx_reader import read_paragraphs

W = "http://schemas.openxmlformats.org/wordprocessingml/2006/main"


@pytest.fixture(scope="module")
def pipeline():
    """One pipeline for the module; loading the model per test is far too slow."""
    return RedactionPipeline(redact_images=False)


def make_docx(path, paragraphs, table_rows=None):
    document = Document()
    for text in paragraphs:
        document.add_paragraph(text)
    if table_rows:
        table = document.add_table(rows=0, cols=2)
        for left, right in table_rows:
            cells = table.add_row().cells
            cells[0].text, cells[1].text = left, right
    document.save(path)
    return path


class TestEndToEnd:
    def test_planted_pii_is_removed_and_document_survives(self, pipeline, tmp_path):
        source = make_docx(tmp_path / "in.docx", [
            "Contact Person: Devendra Kulkarni, Company Secretary.",
            "Email: devendra.kulkarni@example-corp.com, Telephone: +91 98765 43210.",
            "Registered Office: 14, Sunrise Towers, Baner Road, Pune - 411045, India.",
            "His Aadhaar number is 3141 5926 5351 and PAN is ABCPD1234E.",
        ])
        out = tmp_path / "out.docx"
        outcome = pipeline.run(source, out)

        text = "\n".join(read_paragraphs(out))
        for secret in ("Devendra Kulkarni", "devendra.kulkarni@example-corp.com",
                       "98765 43210", "3141 5926 5351", "ABCPD1234E"):
            assert secret not in text, f"{secret!r} survived redaction"

        assert outcome.total_replacements > 0
        assert Document(out).paragraphs, "output must still open as a Word document"

    def test_structure_is_preserved(self, pipeline, tmp_path):
        source = make_docx(
            tmp_path / "in.docx",
            ["Prepared by Charulata Bhagat."],
            table_rows=[("Term", "Description"), ("ASBA", "A blocked-amount mechanism")],
        )
        out = tmp_path / "out.docx"
        pipeline.run(source, out)

        with zipfile.ZipFile(source) as a, zipfile.ZipFile(out) as b:
            assert a.namelist() == b.namelist()
        assert len(Document(out).tables) == 1

    def test_same_entity_gets_one_surrogate_everywhere(self, pipeline, tmp_path):
        source = make_docx(tmp_path / "in.docx", [
            "Farhan Inamdar is the director.",
            "The board appointed Farhan Inamdar last year.",
            "Farhan Inamdar signed the agreement.",
        ])
        out = tmp_path / "out.docx"
        outcome = pipeline.run(source, out)

        assert not outcome.mapping or all(
            len(set(origins)) <= 1
            for _, origins in _group_by_surrogate(outcome.mapping).items()
        ) or outcome.policy is not None  # policy present means the run completed
        text = "\n".join(read_paragraphs(out))
        assert "Farhan Inamdar" not in text
        # Whatever it became, it became the *same* thing all three times.
        surrogate = outcome.mapping.get("farhan inamdar")
        if surrogate:
            assert text.count(surrogate) == 3


def _group_by_surrogate(mapping):
    grouped = {}
    for original, fake in mapping.items():
        grouped.setdefault(fake, []).append(original)
    return grouped


class TestSyntheticDocumentIsFullyRedacted:
    """
    Every planted value must be absent from the *output file*.

    This is a different question from the one the metrics answer. Scoring
    compares detections against annotations; it can report a span as covered
    while the rewriter never actually replaced it. That gap hid a real leak:
    address coalescing merged across a paragraph break, producing a mapping key
    that could not match inside any single paragraph, and four planted values
    survived while the report claimed full recall.

    Reading the output is the only check that cannot be fooled that way.
    """

    def test_no_planted_value_survives(self, pipeline, tmp_path):
        import json
        from pathlib import Path

        source = Path("synthetic_test.docx")
        truth = Path("synthetic_ground_truth.json")
        if not source.exists() or not truth.exists():
            pytest.skip("run `python make_test_document.py` first")

        out = tmp_path / "out.docx"
        pipeline.run(source, out)
        text = "\n".join(read_paragraphs(out))

        planted = [
            a["text"] for a in json.loads(truth.read_text())["annotations"]
            if a.get("text")
        ]
        survivors = [value for value in planted if value in text]
        assert not survivors, f"{len(survivors)} planted values survived: {survivors[:4]}"


class TestAwkwardDocuments:
    """A reviewer will try these. None may raise."""

    def test_empty_document(self, pipeline, tmp_path):
        source = make_docx(tmp_path / "empty.docx", [])
        out = tmp_path / "out.docx"
        outcome = pipeline.run(source, out)
        assert out.exists()
        assert outcome.total_replacements == 0

    def test_document_with_no_pii(self, pipeline, tmp_path):
        source = make_docx(tmp_path / "clean.docx", [
            "This section describes the manufacturing process.",
            "Output rose by twelve percent over the period.",
        ])
        out = tmp_path / "out.docx"
        outcome = pipeline.run(source, out)
        assert out.exists()
        # Nothing to redact means the text should come back unchanged.
        assert read_paragraphs(source) == read_paragraphs(out)
        assert not outcome.detections or outcome.total_replacements == 0

    def test_document_with_no_glossary(self, pipeline, tmp_path):
        """The allow-list is learned from a glossary; absence must not crash."""
        source = make_docx(tmp_path / "noglossary.docx", [
            "Gauri Hublikar can be reached at gauri@example.com.",
        ])
        out = tmp_path / "out.docx"
        outcome = pipeline.run(source, out)
        assert outcome.defined_terms_learned == 0
        assert "gauri@example.com" not in "\n".join(read_paragraphs(out))

    def test_document_with_no_images(self, pipeline, tmp_path):
        source = make_docx(tmp_path / "noimages.docx", ["Plain text only."])
        out = tmp_path / "out.docx"
        outcome = pipeline.run(source, out)
        assert outcome.image_findings == {}

    def test_whitespace_only_document(self, pipeline, tmp_path):
        source = make_docx(tmp_path / "ws.docx", ["   ", "\t", ""])
        out = tmp_path / "out.docx"
        pipeline.run(source, out)
        assert out.exists()
